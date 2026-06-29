import os
import math
import hashlib
import tempfile
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

try:
    from PIL import Image, ImageOps  # type: ignore[import-not-found]
    PIL_DISPONIVEL = True
except ImportError:
    PIL_DISPONIVEL = False

def _resolver_caminho_imagem(base_imagens, imagem):
    """
    Se 'imagem' for caminho absoluto, retorna como está.
    Caso contrário, combina com a pasta base de imagens.
    """
    if os.path.isabs(imagem):
        return imagem
    return os.path.join(base_imagens, imagem)


def _definir_bordas_tabela_brancas(tabela):
    """Define todas as bordas da tabela em branco para ficarem invisíveis."""
    tbl = tabela._tbl
    tbl_pr = tbl.tblPr

    tbl_borders = OxmlElement('w:tblBorders')
    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'FFFFFF')
        tbl_borders.append(border)

    tbl_pr.append(tbl_borders)


def _compactar_imagem(caminho_imagem, pasta_temp, max_lado=1920, qualidade_jpeg=65):
    """
    Compacta a imagem para JPEG, reduzindo resolução e qualidade.
    Retorna o caminho da imagem compactada.
    """
    if not PIL_DISPONIVEL:
        return caminho_imagem

    try:
        sufixo = hashlib.md5(caminho_imagem.encode("utf-8")).hexdigest()[:8]
        nome_saida = f"{os.path.splitext(os.path.basename(caminho_imagem))[0]}_{sufixo}.jpg"
        caminho_saida = os.path.join(pasta_temp, nome_saida)

        with Image.open(caminho_imagem) as img:
            img = ImageOps.exif_transpose(img)
            if img.mode not in ("RGB", "L"):
                img = img.convert("RGB")

            img.thumbnail((max_lado, max_lado), Image.Resampling.LANCZOS)
            img.save(caminho_saida, format="JPEG", quality=qualidade_jpeg, optimize=True, progressive=True)

        return caminho_saida
    except Exception as e:
        print(f"Aviso: Falha ao compactar '{caminho_imagem}': {e}")
        return caminho_imagem


def adicionar_fotos_na_secao(doc, texto_secao, lista_imagens, base_imagens, pasta_temp_compactadas=None, compactar_imagens=True, max_lado=1920, qualidade_jpeg=65):
    """
    Procura por 'texto_secao' no documento e insere as imagens de 2 em 2
    logo após esse parágrafo, adicionando uma quebra de página ao final.
    """
    for paragrafo in doc.paragraphs:
        if texto_secao in paragrafo.text:
            imagens_existentes = []
            for imagem in lista_imagens:
                caminho_imagem = _resolver_caminho_imagem(base_imagens, imagem)
                if not os.path.exists(caminho_imagem):
                    print(f"Aviso: Imagem não encontrada -> {caminho_imagem}")
                    continue
                imagens_existentes.append(caminho_imagem)

            qtd_imagens = len(imagens_existentes)
            if qtd_imagens == 0:
                print(f"Aviso: Nenhuma imagem válida para '{texto_secao}'.")
                return

            # Calcula a quantidade de linhas necessárias (teto da divisão por 2)
            linhas = math.ceil(qtd_imagens / 2)
            
            # Cria a tabela no documento
            tabela = doc.add_table(rows=linhas, cols=2)
            tabela.autofit = False
            _definir_bordas_tabela_brancas(tabela)
            
            # Move a tabela para logo abaixo do parágrafo do "Relatório Fotográfico"
            paragrafo._p.addnext(tabela._tbl)
            
            # Insere as imagens nas células da tabela
            for i, caminho_imagem in enumerate(imagens_existentes):
                caminho_para_inserir = caminho_imagem
                if compactar_imagens and pasta_temp_compactadas:
                    caminho_para_inserir = _compactar_imagem(
                        caminho_imagem,
                        pasta_temp_compactadas,
                        max_lado=max_lado,
                        qualidade_jpeg=qualidade_jpeg,
                    )
                    
                linha_atual = i // 2
                coluna_atual = i % 2
                celula = tabela.cell(linha_atual, coluna_atual)
                
                # Adiciona a imagem no parágrafo dentro da célula
                p_celula = celula.paragraphs[0]
                p_celula.alignment = 1 # Centraliza a imagem na célula
                run = p_celula.add_run()
                
                # Ajuste a largura (Inches) conforme o tamanho ideal para sua página
                run.add_picture(caminho_para_inserir, width=Inches(3.0))

                # Adiciona legenda editável abaixo da foto
                #legenda = os.path.splitext(os.path.basename(caminho_imagem))[0]
                #p_legenda = celula.add_paragraph(f"Legenda: {legenda}")
                #p_legenda.alignment = 1

            # Adiciona uma quebra de página logo após a tabela
            # Cria um novo elemento XML de parágrafo e insere a quebra
            p_quebra = doc.add_paragraph()
            p_quebra.add_run().add_break(WD_BREAK.PAGE)
            tabela._tbl.addnext(p_quebra._p)
            
            print(f"Sucesso: {qtd_imagens} imagens inseridas em '{texto_secao}'.")
            break # Interrompe após encontrar e processar a seção

def main():
    # 1. Caminho do seu documento base (template)
    caminho_documento = "/var/home/gladistony/.var/app/org.telegram.desktop/data/TelegramDesktop/tdata/temp_data/01_MODELO_RELATORIO_VISTORIA.docx"
    doc = Document(caminho_documento)

    # 2. Pasta padrão absoluta onde estão as imagens
    pasta_imagens = "/home/gladistony/Documentos/BRADESCO/"

    # 3. Dicionário mapeando a seção para a lista de imagens desejadas
    # Aqui você pode informar somente os nomes dos arquivos
    imagens_por_secao = {
        "1.4. Relatório Fotográfico": [f"extintor{i}.jpg" for i in range(1, 51)],
        "2.4. Relatório Fotográfico": [f"luz{i}.jpg" for i in range(1, 13)],
        "3.4. Relatório Fotográfico": [f"alarme{i}.jpg" for i in range(1, 16)],
        "5.3. Relatório Fotográfico": [f"corrimão{i}.jpg" for i in range(1, 10)],
        "6.1.2. Relatório Fotográfico": [f"hidrante{i}.jpg" for i in range(1, 7)],
        "6.3.2. Relatório Fotográfico": [f"mangue{i}.jpg" for i in range(1, 4)],
        "7.3. Relatório Fotográfico": [f"saida{i}.jpg" for i in range(1, 16)],
        "9.3. Relatório Fotográfico": [f"dijuntor{i}.jpg" for i in range(1, 14)],
        # Adicione as demais seções conforme necessário...
    }
    imagens_por_secao["3.4. Relatório Fotográfico"].extend([f"detector{i}.jpg" for i in range(1, 32)])

    # 4. Configuração de compactação (reduz bastante o tamanho do .docx)
    compactar_imagens = True
    max_lado = 1920      # px (ex.: 1600, 1920, 2560)
    qualidade_jpeg = 60  # 1 a 95 (quanto menor, menor o arquivo)

    if compactar_imagens and not PIL_DISPONIVEL:
        print("Aviso: Pillow não está instalado. Rode: pip install Pillow")
        print("Continuando sem compactação de imagens...")

    # 5. Processa cada seção definida no dicionário
    with tempfile.TemporaryDirectory(prefix="img_relatorio_") as pasta_temp_compactadas:
        for secao, imagens in imagens_por_secao.items():
            adicionar_fotos_na_secao(
                doc,
                secao,
                imagens,
                pasta_imagens,
                pasta_temp_compactadas=pasta_temp_compactadas,
                compactar_imagens=compactar_imagens and PIL_DISPONIVEL,
                max_lado=max_lado,
                qualidade_jpeg=qualidade_jpeg,
            )

    # 6. Salva o novo documento preenchido
    nome_arquivo_saida = "RELATORIO_VISTORIA_PREENCHIDO.docx"
    doc.save(nome_arquivo_saida)
    print(f"\nConcluído! Documento salvo como: {nome_arquivo_saida}")

if __name__ == "__main__":
    main()